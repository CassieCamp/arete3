import os
import logging
from typing import Optional
import pypdf
from docx import Document as DocxDocument
import io
import tempfile
from fastapi import UploadFile
import uuid
from datetime import datetime

logger = logging.getLogger(__name__)


async def save_uploaded_file(file: UploadFile) -> str:
    """
    Save an uploaded file to the uploads directory and return the file path.
    
    Args:
        file: FastAPI UploadFile object
        
    Returns:
        str: Path to the saved file
        
    Raises:
        Exception: If file saving fails
    """
    logger.info(f"=== save_uploaded_file called ===")
    logger.info(f"filename: {file.filename}, content_type: {file.content_type}")
    
    try:
        # Create uploads directory if it doesn't exist
        uploads_dir = "uploads"
        os.makedirs(uploads_dir, exist_ok=True)
        
        # Generate unique filename to avoid conflicts
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        
        # Get file extension
        original_filename = file.filename or "uploaded_file"
        file_extension = ""
        if "." in original_filename:
            file_extension = "." + original_filename.split(".")[-1]
        
        # Create unique filename
        safe_filename = f"{timestamp}_{unique_id}_{original_filename.replace(' ', '_')}"
        if not safe_filename.endswith(file_extension):
            safe_filename += file_extension
            
        file_path = os.path.join(uploads_dir, safe_filename)
        
        # Read file content
        file_content = await file.read()
        
        # Save file to disk
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        logger.info(f"✅ Successfully saved file to: {file_path}")
        return file_path
        
    except Exception as e:
        error_msg = f"Error saving uploaded file: {str(e)}"
        logger.error(error_msg)
        raise Exception(error_msg) from e


async def extract_text_from_file(file_path: str, content_type: str) -> str:
    """
    Extract text content from a file based on its content type.
    
    Args:
        file_path: Path to the file to extract text from
        content_type: MIME type of the file
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
        Exception: For other processing errors
    """
    logger.info(f"=== extract_text_from_file called ===")
    logger.info(f"file_path: {file_path}, content_type: {content_type}")
    
    # Check if file exists
    if not os.path.exists(file_path):
        error_msg = f"File not found: {file_path}"
        logger.error(error_msg)
        raise FileNotFoundError(error_msg)
    
    # Map content types to extraction methods
    content_type_mapping = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'docx',  # Treat .doc as .docx (limited support)
        'text/plain': 'txt'
    }
    
    # Determine file type from content type
    file_type = content_type_mapping.get(content_type)
    
    # If content type mapping fails, try to determine from file extension
    if not file_type:
        file_extension = file_path.lower().split('.')[-1] if '.' in file_path else 'txt'
        extension_mapping = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',
            'txt': 'txt'
        }
        file_type = extension_mapping.get(file_extension, 'txt')
    
    # Use the existing TextExtractionService
    text_service = TextExtractionService()
    
    if not text_service.is_supported_file_type(file_type):
        error_msg = f"Unsupported file type: {file_type}. Supported types: {text_service.get_supported_file_types()}"
        logger.error(error_msg)
        raise ValueError(error_msg)
    
    try:
        extracted_text = text_service.extract_text_from_file(file_path, file_type)
        logger.info(f"✅ Successfully extracted {len(extracted_text)} characters from {file_path}")
        return extracted_text
        
    except Exception as e:
        error_msg = f"Error extracting text from {file_path}: {str(e)}"
        logger.error(error_msg)
        raise Exception(error_msg) from e


class TextExtractionService:
    """Service for extracting text content from various document types"""
    
    def __init__(self):
        self.supported_types = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt
        }
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from a file based on its type.
        
        Args:
            file_path: Path to the file to extract text from
            file_type: Type of the file (pdf, docx, txt, etc.)
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
            Exception: For other processing errors
        """
        logger.info(f"=== TextExtractionService.extract_text_from_file called ===")
        logger.info(f"file_path: {file_path}, file_type: {file_type}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        # Normalize file type to lowercase
        file_type = file_type.lower()
        
        # Check if file type is supported
        if file_type not in self.supported_types:
            error_msg = f"Unsupported file type: {file_type}. Supported types: {list(self.supported_types.keys())}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        try:
            # Extract text using the appropriate method
            extraction_method = self.supported_types[file_type]
            extracted_text = extraction_method(file_path)
            
            logger.info(f"✅ Successfully extracted {len(extracted_text)} characters from {file_path}")
            return extracted_text
            
        except Exception as e:
            error_msg = f"Error extracting text from {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.error(f"Exception type: {type(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise Exception(error_msg) from e
    
    async def extract_text_from_bytes(self, file_content: bytes, filename: str) -> str:
        """
        Extract text content from file bytes based on filename extension.
        
        Args:
            file_content: Raw bytes of the file
            filename: Name of the file (used to determine file type)
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            Exception: For other processing errors
        """
        logger.info(f"=== TextExtractionService.extract_text_from_bytes called ===")
        logger.info(f"filename: {filename}, content_size: {len(file_content)} bytes")
        
        # Determine file type from filename extension
        file_extension = filename.lower().split('.')[-1] if '.' in filename else 'txt'
        
        # Map common extensions to our supported types
        extension_mapping = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',  # Treat .doc as .docx (limited support)
            'txt': 'txt',
            'text': 'txt'
        }
        
        file_type = extension_mapping.get(file_extension, 'txt')
        
        # Check if file type is supported
        if file_type not in self.supported_types:
            error_msg = f"Unsupported file type: {file_type}. Supported types: {list(self.supported_types.keys())}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        try:
            # Create a temporary file to work with existing extraction methods
            with tempfile.NamedTemporaryFile(suffix=f'.{file_extension}', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text using the appropriate method
                extraction_method = self.supported_types[file_type]
                extracted_text = extraction_method(temp_file_path)
                
                logger.info(f"✅ Successfully extracted {len(extracted_text)} characters from {filename}")
                return extracted_text
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temporary file {temp_file_path}: {cleanup_error}")
                    
        except Exception as e:
            error_msg = f"Error extracting text from {filename}: {str(e)}"
            logger.error(error_msg)
            logger.error(f"Exception type: {type(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise Exception(error_msg) from e
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyPDF2"""
        logger.info(f"Extracting text from PDF: {file_path}")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text_content.append(page_text)
                            logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                    except Exception as page_error:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                        continue
            
            extracted_text = '\n\n'.join(text_content)
            
            if not extracted_text.strip():
                logger.warning(f"No text content extracted from PDF: {file_path}")
                return "No text content could be extracted from this PDF file."
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx"""
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n'.join(text_content)
            
            if not extracted_text.strip():
                logger.warning(f"No text content extracted from DOCX: {file_path}")
                return "No text content could be extracted from this DOCX file."
            
            logger.info(f"Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        logger.info(f"Extracting text from TXT: {file_path}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        logger.info(f"Successfully read TXT file with {encoding} encoding")
                        return content
                except UnicodeDecodeError:
                    logger.debug(f"Failed to read with {encoding} encoding, trying next...")
                    continue
            
            # If all encodings fail, try reading as binary and decode with errors='replace'
            logger.warning(f"All standard encodings failed for {file_path}, using fallback method")
            with open(file_path, 'rb') as file:
                content = file.read().decode('utf-8', errors='replace')
                return content
                
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise
    
    def is_supported_file_type(self, file_type: str) -> bool:
        """Check if a file type is supported for text extraction"""
        return file_type.lower() in self.supported_types
    
    def get_supported_file_types(self) -> list:
        """Get list of supported file types"""
        return list(self.supported_types.keys())